import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Inches
import re
import spacy
from datetime import datetime, timedelta
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors
from io import BytesIO
import sqlite3
import json
from functools import lru_cache
import subprocess
import sys

# ========== CONFIGURACIÓN DE PÁGINA ==========
st.set_page_config(page_title="Contract AI Assistant - Professional", layout="wide", page_icon="⚖️", initial_sidebar_state="expanded")

# ========== TEMA OSCURO/CLARO ==========
if "theme" not in st.session_state:
    st.session_state.theme = "light"
def toggle_theme():
    st.session_state.theme = "dark" if st.session_state.theme == "light" else "light"
st.sidebar.button("🌓 Cambiar tema", on_click=toggle_theme)

if st.session_state.theme == "dark":
    st.markdown("""
    <style>
    .stApp { background-color: #0E1117; color: #FAFAFA; }
    .css-1d391kg { background-color: #1E1E1E; }
    .risk-low { background-color: #22C55E; padding: 0.2rem 1rem; border-radius: 20px; color: black; display: inline-block; }
    .risk-medium { background-color: #EAB308; padding: 0.2rem 1rem; border-radius: 20px; color: black; display: inline-block; }
    .risk-high { background-color: #EF4444; padding: 0.2rem 1rem; border-radius: 20px; color: white; display: inline-block; }
    </style>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <style>
    .risk-low { background-color: #22C55E; padding: 0.2rem 1rem; border-radius: 20px; color: white; display: inline-block; }
    .risk-medium { background-color: #EAB308; padding: 0.2rem 1rem; border-radius: 20px; color: white; display: inline-block; }
    .risk-high { background-color: #EF4444; padding: 0.2rem 1rem; border-radius: 20px; color: white; display: inline-block; }
    </style>
    """, unsafe_allow_html=True)

st.markdown('<h1 style="text-align: center;">⚖️ Contract AI Assistant <span style="font-size: 0.6rem;">Professional</span></h1>', unsafe_allow_html=True)
st.markdown("---")

# ========== DESCARGA AUTOMÁTICA DEL MODELO spaCY ==========
@st.cache_resource
def load_nlp():
    try:
        return spacy.load("es_core_news_sm")
    except OSError:
        with st.spinner("Descargando modelo de lenguaje español (solo la primera vez)... puede tomar un minuto."):
            subprocess.run([sys.executable, "-m", "spacy", "download", "es_core_news_sm"])
            return spacy.load("es_core_news_sm")
    except Exception as e:
        st.error(f"Error al cargar el modelo de lenguaje: {e}")
        return None

nlp = load_nlp()

# ========== BASE DE DATOS (HISTORIAL) CON MIGRACIÓN ==========
DB_PATH = "contract_ai.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS history
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  filename TEXT,
                  date TEXT,
                  summary TEXT,
                  risk_level TEXT,
                  risk_score REAL,
                  parties TEXT,
                  key_dates TEXT,
                  clauses TEXT)''')
    c.execute("PRAGMA table_info(history)")
    columns = [col[1] for col in c.fetchall()]
    if 'risk_score' not in columns:
        c.execute("ALTER TABLE history ADD COLUMN risk_score REAL DEFAULT 0")
    if 'clauses' not in columns:
        c.execute("ALTER TABLE history ADD COLUMN clauses TEXT DEFAULT '[]'")
    conn.commit()
    conn.close()

init_db()

def clear_history():
    """Elimina todos los registros del historial."""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("DELETE FROM history")
    conn.commit()
    conn.close()
    st.session_state.history_cleared = True

def save_analysis(filename, summary, risk_level, risk_score, parties, key_dates, clauses):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("INSERT INTO history (filename, date, summary, risk_level, risk_score, parties, key_dates, clauses) VALUES (?,?,?,?,?,?,?,?)",
              (filename, datetime.now().isoformat(), summary[:500], risk_level, risk_score, json.dumps(parties), json.dumps(key_dates), json.dumps(clauses)))
    conn.commit()
    conn.close()

def load_history():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, filename, date, risk_level, risk_score FROM history ORDER BY date DESC LIMIT 20")
    data = c.fetchall()
    conn.close()
    return data

def load_full_analysis(analysis_id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT summary, risk_level, risk_score, parties, key_dates, clauses FROM history WHERE id=?", (analysis_id,))
    row = c.fetchone()
    conn.close()
    if row:
        return {
            "summary": row[0],
            "risk_level": row[1],
            "risk_score": row[2],
            "parties": json.loads(row[3]),
            "key_dates": json.loads(row[4]),
            "clauses": json.loads(row[5])
        }
    return None

# ========== FUNCIONES AUXILIARES ==========
STOP_ENTITIES = {
    "entregar", "mantener", "pagar", "proporcionar", "obligaciones", "objeto", "monto",
    "cliente", "prestador", "parte", "contrato", "fecha", "lugar", "firma", "nombre",
    "domicilio", "comparece", "actuando", "mayor", "edad", "cedula", "ciudadania",
    "se", "obliga", "entre", "prestacion", "servicios", "profesionales", "el", "la", "los"
}

def extract_text(file):
    if file.type == "application/pdf":
        try:
            with pdfplumber.open(file) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            return text.strip()
        except:
            return None
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = Document(file)
            return "\n".join([para.text for para in doc.paragraphs]).strip()
        except:
            return None
    else:
        try:
            return file.read().decode("utf-8").strip()
        except:
            return None

def extract_parties(text):
    parties = []
    pattern1 = r'(?:DE UNA PARTE|DE OTRA PARTE|CONTRATANTE|CONTRATISTA|PRESTADOR|CLIENTE)[:\s]+([A-Z][a-záéíóúñ]+(?:\s+[A-Z][a-záéíóúñ]+)*)'
    matches = re.findall(pattern1, text, re.IGNORECASE)
    parties.extend(matches)
    pattern2 = r'(?:Nombre|Nombres?)[:\s]+([A-Z][a-záéíóúñ]+(?:\s+[A-Z][a-záéíóúñ]+)*)'
    matches = re.findall(pattern2, text, re.IGNORECASE)
    parties.extend(matches)
    pattern3 = r'([A-Z][a-záéíóúñ]+(?:\s+[A-Z][a-záéíóúñ]+){1,3})\s*(?:,|,?\s+mayor de edad|,?\s+con cédula)'
    matches = re.findall(pattern3, text)
    parties.extend(matches)
    valid = []
    for p in parties:
        p = p.strip()
        if len(p.split()) >= 2 and p.lower() not in STOP_ENTITIES:
            if not any(word in p.lower() for word in ["se obliga", "entre", "prestacion"]):
                valid.append(p)
    return list(set(valid))

def extract_dates(text):
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
        r'\b\d{1,2}\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+\d{4}\b',
        r'\b(?:vigencia|término|vencimiento|fecha de inicio|fecha de finalización)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        r'\b(?:a partir del|desde el|hasta el)\s+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
    ]
    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        dates.extend(matches)
    cleaned = []
    for d in set(dates):
        try:
            if '-' in d:
                parts = d.split('-')
            elif '/' in d:
                parts = d.split('/')
            else:
                continue
            if len(parts) == 3:
                day, month, year = parts
                cleaned.append(datetime(int(year), int(month), int(day)))
        except:
            pass
    return cleaned

def extract_money(text):
    patterns = [
        r'\$\s*[\d,]+(?:\.\d{2})?',
        r'(\d+(?:\.\d{3})*(?:,\d{2})?)\s*(?:USD|EUR|COP)'
    ]
    amounts = []
    for p in patterns:
        matches = re.findall(p, text)
        amounts.extend(matches)
    numeric = []
    for a in amounts:
        try:
            num = float(re.sub(r'[^\d.,]', '', a).replace(',', ''))
            numeric.append(num)
        except:
            pass
    return numeric

def analyze_contract(text):
    if not text or not nlp:
        return None, None
    doc = nlp(text[:50000])
    
    entities = {"PER": [], "ORG": [], "DATE": [], "MONEY": []}
    for ent in doc.ents:
        if ent.label_ in entities and ent.text[0].isupper() and len(ent.text) > 2:
            entities[ent.label_].append(ent.text)
    
    parties = extract_parties(text)
    entities["PER"].extend(parties)
    dates = extract_dates(text)
    entities["DATE"].extend([d.strftime("%Y-%m-%d") for d in dates])
    amounts = extract_money(text)
    entities["MONEY"].extend([str(round(a,2)) for a in amounts])
    
    for k in entities:
        entities[k] = list(set(entities[k]))[:5]
    
    risk_keywords = {
        "confidencialidad": r"confidencial|secreto|no divulgación",
        "indemnización": r"indemniz[aá]|responsabilidad|daños y perjuicios",
        "resolución anticipada": r"resolución anticipada|terminación adelantada|incumplimiento grave",
        "penalización": r"penalizaci[oó]n|multa|interés moratorio",
        "jurisdicción": r"jurisdicci[oó]n|arbitraje|tribunal|ley aplicable",
        "garantía": r"garant[íi]a|aval|fianza",
        "cesión": r"cesi[oó]n|transferencia|subcontratación",
        "renovación automática": r"renovación automática|prórroga automática",
        "exclusividad": r"exclusividad|exclusivo",
        "no competencia": r"no competencia|no competir",
        "propiedad intelectual": r"propiedad intelectual|derechos de autor|patente",
        "fuerza mayor": r"fuerza mayor|caso fortuito",
        "rescisión": r"rescisión|rescindir|dar por terminado",
        "plazo": r"plazo\s+(?:de\s+)?(\d+\s*(?:días|meses|años))",
        "precio": r"precio\s+(?:total|mensual|anual)[:\s]*\$?\s*[\d,]+"
    }
    
    paragraphs = re.split(r'\n\s*\n|\.\s+', text)
    risk_clauses = {}
    for clause, pattern in risk_keywords.items():
        contexts = []
        for para in paragraphs:
            if re.search(pattern, para, re.IGNORECASE):
                clean = re.sub(r'\s+', ' ', para).strip()
                contexts.append(clean[:250])
                if len(contexts) >= 2:
                    break
        if contexts:
            risk_clauses[clause] = contexts
    
    risk_score = min(100, len(risk_clauses) * 12)
    if risk_score >= 60:
        risk_level = "alto"
    elif risk_score >= 30:
        risk_level = "medio"
    else:
        risk_level = "bajo"
    
    summary = "**Resumen Ejecutivo**\n\n"
    if entities["PER"]:
        summary += f"👥 **Partes involucradas:** {', '.join(entities['PER'][:5])}\n\n"
    if entities["ORG"]:
        summary += f"🏢 **Organizaciones:** {', '.join(entities['ORG'][:5])}\n\n"
    if entities["DATE"]:
        summary += f"📅 **Fechas relevantes:** {', '.join(entities['DATE'][:5])}\n\n"
    if entities["MONEY"]:
        summary += f"💰 **Montos detectados:** {', '.join(entities['MONEY'][:3])}\n\n"
    risk_emoji = "🟢" if risk_level == "bajo" else "🟡" if risk_level == "medio" else "🔴"
    summary += f"{risk_emoji} **Nivel de riesgo:** {risk_level.upper()} (puntuación: {risk_score}/100)\n\n"
    if risk_clauses:
        summary += "**⚖️ Cláusulas críticas detectadas:**\n"
        for clause, ctx in risk_clauses.items():
            summary += f"- **{clause.capitalize()}**: {ctx[0][:120]}...\n"
    else:
        summary += "No se detectaron cláusulas de alto riesgo.\n"
    
    return summary, {"entities": entities, "risk_clauses": risk_clauses, "risk_level": risk_level, "risk_score": risk_score, "dates": dates}

# ========== FUNCIONES DE EXPORTACIÓN ==========
def generate_pdf(summary, metadata, details):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Title'], alignment=TA_CENTER, fontSize=16)
    normal_style = styles['Normal']
    
    clean = re.sub(r'[^\x00-\x7F\u00C0-\u00FF]+', '', summary)
    clean = re.sub(r'\*\*', '', clean)
    clean = clean.replace('\n\n', '<br/><br/>')
    
    story = []
    story.append(Paragraph("Contract AI Assistant - Informe Legal", title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(clean, normal_style))
    story.append(Spacer(1, 12))
    
    data = [
        ["Nivel de riesgo", details["risk_level"].upper()],
        ["Puntuación de riesgo", f"{details['risk_score']}/100"],
        ["Cláusulas críticas", str(len(details["risk_clauses"]))],
        ["Partes detectadas", str(len(details["entities"].get("PER", [])))],
        ["Fechas encontradas", str(len(details.get("dates", [])))]
    ]
    t = Table(data, colWidths=[100, 100])
    t.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, colors.grey), ('BACKGROUND', (0,0), (0,-1), colors.lightgrey)]))
    story.append(t)
    story.append(Spacer(1, 12))
    
    story.append(Paragraph(f"Fecha análisis: {metadata['date'][:19]}", normal_style))
    story.append(Paragraph(f"Documento: {metadata['filename']}", normal_style))
    doc.build(story)
    return buffer.getvalue()

def generate_docx(summary, metadata, details):
    doc = Document()
    doc.add_heading('Contract AI Assistant - Informe Legal', 0)
    doc.add_paragraph(summary.replace('**', ''))
    doc.add_paragraph(f"\nMétricas:")
    doc.add_paragraph(f"- Nivel de riesgo: {details['risk_level'].upper()}")
    doc.add_paragraph(f"- Puntuación de riesgo: {details['risk_score']}/100")
    doc.add_paragraph(f"- Cláusulas críticas: {len(details['risk_clauses'])}")
    doc.add_paragraph(f"- Fechas encontradas: {len(details.get('dates', []))}")
    doc.add_paragraph(f"\nFecha análisis: {metadata['date'][:19]}")
    doc.add_paragraph(f"Documento: {metadata['filename']}")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ========== INTERFAZ PRINCIPAL ==========
with st.sidebar:
    st.header("📁 Historial")
    history = load_history()
    if history:
        for h in history:
            st.write(f"📄 {h[1]} - {h[2][:10]} ({h[3]}, {h[4]:.0f}/100)")
        if st.button("🗑️ Limpiar historial"):
            clear_history()
            st.rerun()
    else:
        st.info("Aún no hay análisis.")
    st.markdown("---")
    st.caption("Contract AI Assistant Professional | NLP con spaCy")

uploaded = st.file_uploader("📂 Sube tu contrato (PDF, DOCX o TXT)", type=["pdf", "docx", "txt"])

if uploaded:
    filename = uploaded.name
    with st.spinner("Extrayendo texto..."):
        text = extract_text(uploaded)
    if not text:
        st.error("No se pudo extraer texto.")
        st.stop()
    
    with st.expander("📄 Vista previa del texto extraído"):
        st.text_area("", text[:2000], height=200)
    
    if st.button("🔍 Analizar contrato", type="primary"):
        with st.spinner("Analizando con IA..."):
            analysis, details = analyze_contract(text)
        if analysis:
            parties = details["entities"].get("PER", [])
            dates_str = [d.strftime("%Y-%m-%d") for d in details.get("dates", [])]
            save_analysis(filename, analysis, details["risk_level"], details["risk_score"], parties, dates_str, list(details["risk_clauses"].keys()))
            
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Resumen", "⚖️ Cláusulas", "📅 Fechas y montos", "📊 Dashboard", "📈 Comparar"])
            with tab1:
                st.markdown(analysis, unsafe_allow_html=True)
            with tab2:
                if details["risk_clauses"]:
                    for clause, contexts in details["risk_clauses"].items():
                        with st.expander(f"🔍 {clause.capitalize()}"):
                            st.write(contexts[0])
                else:
                    st.info("No se detectaron cláusulas críticas.")
            with tab3:
                col1, col2 = st.columns(2)
                with col1:
                    st.subheader("📅 Fechas")
                    if details.get("dates"):
                        for d in details["dates"]:
                            st.write(f"- {d.strftime('%d/%m/%Y')}")
                    else:
                        st.write("No se encontraron fechas.")
                with col2:
                    st.subheader("💰 Montos")
                    amounts = details["entities"].get("MONEY", [])
                    if amounts:
                        for a in amounts:
                            st.write(f"- {a}")
                    else:
                        st.write("No se detectaron montos.")
            with tab4:
                categories = list(details["risk_clauses"].keys())
                if categories:
                    values = [1] * len(categories)
                    fig = go.Figure(data=[go.Bar(x=categories[:5], y=values[:5], marker_color='indianred', text=values, textposition='auto')])
                    fig.update_layout(title="Cláusulas detectadas (primeras 5)", xaxis_title="Tipo", yaxis_title="Presencia", height=400)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("Sin cláusulas de riesgo para graficar.")
                fig2 = go.Figure(go.Indicator(
                    mode="gauge+number+delta",
                    value=details["risk_score"],
                    title={"text": "Riesgo global"},
                    delta={"reference": 50},
                    gauge={"axis": {"range": [0,100]}, "bar": {"color": "green" if details["risk_score"]<30 else "orange" if details["risk_score"]<60 else "red"},
                           "steps": [{"range": [0,33], "color": "lightgreen"}, {"range": [33,66], "color": "yellow"}, {"range": [66,100], "color": "red"}]}
                ))
                fig2.update_layout(height=300)
                st.plotly_chart(fig2, use_container_width=True)
                if details.get("dates"):
                    df_dates = pd.DataFrame({"Fecha": details["dates"]})
                    fig3 = px.scatter(df_dates, x="Fecha", title="Línea de tiempo de fechas")
                    st.plotly_chart(fig3, use_container_width=True)
            with tab5:
                st.subheader("Comparador de contratos")
                colC1, colC2 = st.columns(2)
                with colC1:
                    file1 = st.file_uploader("Contrato A", type=["pdf","docx","txt"], key="comp1")
                with colC2:
                    file2 = st.file_uploader("Contrato B", type=["pdf","docx","txt"], key="comp2")
                if file1 and file2 and st.button("Comparar ahora"):
                    text1 = extract_text(file1)
                    text2 = extract_text(file2)
                    if text1 and text2:
                        _, det1 = analyze_contract(text1)
                        _, det2 = analyze_contract(text2)
                        st.subheader("📊 Diferencias clave")
                        colR1, colR2 = st.columns(2)
                        with colR1:
                            st.markdown(f"**{file1.name}**")
                            st.write(f"Partes: {', '.join(det1['entities'].get('PER',[])[:3])}")
                            st.write(f"Riesgo: {det1['risk_level']} ({det1['risk_score']}/100)")
                            st.write(f"Cláusulas: {len(det1['risk_clauses'])}")
                        with colR2:
                            st.markdown(f"**{file2.name}**")
                            st.write(f"Partes: {', '.join(det2['entities'].get('PER',[])[:3])}")
                            st.write(f"Riesgo: {det2['risk_level']} ({det2['risk_score']}/100)")
                            st.write(f"Cláusulas: {len(det2['risk_clauses'])}")
                    else:
                        st.error("Error al extraer texto.")
            
            st.markdown("---")
            st.subheader("📥 Exportar resultados")
            colD1, colD2, colD3 = st.columns(3)
            with colD1:
                st.download_button("📄 TXT", analysis, f"analisis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", "text/plain")
            with colD2:
                pdf_data = generate_pdf(analysis, {"filename": filename, "date": datetime.now().isoformat()}, details)
                st.download_button("📑 PDF", pdf_data, f"reporte_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", "application/pdf")
            with colD3:
                docx_data = generate_docx(analysis, {"filename": filename, "date": datetime.now().isoformat()}, details)
                st.download_button("📄 Word", docx_data, f"reporte_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.error("Error en el análisis.")

# Footer
st.markdown("---")
st.caption("Contract AI Assistant Professional | Análisis legal con IA | Sin licencias | Todos los derechos reservados")
